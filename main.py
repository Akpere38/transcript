import os
import random
from flask import Flask, abort, render_template, redirect, url_for, flash, request, session, send_file
from flask_wtf import FlaskForm
from wtforms import StringField, SubmitField, PasswordField, EmailField, IntegerField, SelectField, DateField, \
    BooleanField, FileField, TimeField
from wtforms.validators import DataRequired, URL, Email
from docx import Document
from flask_bootstrap import Bootstrap5

app = Flask(__name__)
Bootstrap5(app)
app.secret_key = 'your_secret_key'
app.config['UPLOAD_FOLDER'] = 'uploads'


class FileForm(FlaskForm):
    file = FileField("Upload your file(.docx file only)")
    submit = SubmitField("Submit")




@app.route("/", methods=['GET', 'POST'])
def home():
    form = FileForm()
    if form.validate_on_submit():
        file = form.file.data
        if file:
            # print(file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)

            doc = Document(file_path)
            combine_text = ""

            # Iterate through the paragraphs in the document and print the text
            for paragraph in doc.paragraphs:
                combine_text += paragraph.text + ""
            # print(combine_text)
            numeric = ""
            for n in range(5):
                numeric += str(random.randint(0, 10))



            new_file = Document()
            new_file.add_paragraph(combine_text)
            filename = f"transcribe{numeric}.docx"
            new_file.save(f"downloads/{filename}")

            # print(new_file.filename)
            return redirect((url_for("download", filename=filename)))



    return render_template("index.html", form=form)


@app.route("/downloads/<filename>")
def download(filename):
    file_path = f"downloads/{filename}"
    return send_file(path_or_file=file_path, as_attachment=True)



if __name__ == '__main__':
    app.run()
